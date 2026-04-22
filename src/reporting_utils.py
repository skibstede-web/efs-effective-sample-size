"""Reusable reporting helpers for notebook exports."""

from __future__ import annotations

from datetime import datetime
import json
from pathlib import Path
import re
import subprocess
from types import SimpleNamespace
from typing import Any

import matplotlib.pyplot as plt
import pandas as pd
from matplotlib.figure import Figure


def create_run_directory(project_root: Path | str, setup_name: str) -> tuple[str, Path]:
    """Create a timestamped report directory for a setup."""
    root = Path(project_root)
    if not setup_name or not isinstance(setup_name, str):
        raise ValueError("setup_name must be a non-empty string.")
    if not root.exists():
        raise FileNotFoundError(f"Project root does not exist: {root}")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    run_dir = root / "reports" / setup_name / timestamp
    run_dir.mkdir(parents=True, exist_ok=False)
    return timestamp, run_dir


def save_dataframe_csv(df: pd.DataFrame | None, output_path: Path | str) -> Path | None:
    """Save a DataFrame to CSV if present."""
    if df is None:
        return None
    if not isinstance(df, pd.DataFrame):
        raise TypeError("df must be a pandas DataFrame or None.")
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(output, index=False, encoding="utf-8")
    return output


def save_dataframes_dict(
    dataframes_dict: dict[str, pd.DataFrame | None],
    run_dir: Path | str,
) -> list[str]:
    """Save a mapping of file stems to DataFrames under a run directory."""
    if not isinstance(dataframes_dict, dict):
        raise TypeError("dataframes_dict must be a dictionary.")
    run_path = Path(run_dir)
    saved_files: list[str] = []
    for name, df in dataframes_dict.items():
        if not name or df is None:
            continue
        filename = name if name.lower().endswith(".csv") else f"{name}.csv"
        saved_path = save_dataframe_csv(df, run_path / filename)
        if saved_path is not None:
            saved_files.append(saved_path.name)
    return saved_files


def save_figure(fig: Figure | None, output_path: Path | str, dpi: int = 300) -> Path | None:
    """Save a matplotlib figure to PNG if present."""
    if fig is None:
        return None
    if not isinstance(fig, Figure):
        raise TypeError("fig must be a matplotlib.figure.Figure or None.")
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(output, dpi=dpi, bbox_inches="tight")
    return output


def save_figures_dict(
    figures_dict: dict[str, Figure | None],
    run_dir: Path | str,
    dpi: int = 300,
) -> list[str]:
    """Save a mapping of file stems to PNG figures under a run directory."""
    if not isinstance(figures_dict, dict):
        raise TypeError("figures_dict must be a dictionary.")
    run_path = Path(run_dir)
    saved_files: list[str] = []
    for name, fig in figures_dict.items():
        if not name or fig is None:
            continue
        filename = name if name.lower().endswith(".png") else f"{name}.png"
        saved_path = save_figure(fig, run_path / filename, dpi=dpi)
        if saved_path is not None:
            saved_files.append(saved_path.name)
    return saved_files


def write_metadata_json(metadata: dict[str, Any], output_path: Path | str) -> Path:
    """Write metadata as UTF-8 JSON."""
    if not isinstance(metadata, dict):
        raise TypeError("metadata must be a dictionary.")
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(
        json.dumps(metadata, indent=2, ensure_ascii=False, default=_json_default),
        encoding="utf-8",
    )
    return output


def looks_like_python_code(text: str) -> bool:
    """Return True when markdown text strongly resembles pasted Python code."""
    normalized = (text or "").replace("\r\n", "\n")
    stripped_lines = [line.rstrip() for line in normalized.split("\n")]

    # Strip out code-fenced blocks before analysis — they often contain
    # ASCII-art diagrams that look like indented code.
    outside_fence_lines: list[str] = []
    in_fence = False
    for line in stripped_lines:
        if line.strip().startswith("```"):
            in_fence = not in_fence
            continue
        if not in_fence:
            outside_fence_lines.append(line)

    content_lines = [line for line in outside_fence_lines if line.strip()]
    if not content_lines:
        return False

    # If the cell has strong markdown signals, it is almost certainly
    # narrative text, not pasted code.
    markdown_signals = 0
    markdown_signals += sum(1 for line in content_lines if re.match(r"^#{1,6}\s", line.strip()))
    markdown_signals += sum(1 for line in content_lines if line.strip().startswith("|") and line.strip().endswith("|"))
    markdown_signals += sum(1 for line in content_lines if "$$" in line)
    markdown_signals += sum(1 for line in content_lines if re.search(r"\*\*[^*]+\*\*", line))
    if markdown_signals >= 3:
        return False

    code_like_signals = 0

    strong_substrings = (
        "plt.",
        "ax.",
        "fig.",
        "np.",
        "pd.",
        "tight_layout(",
        ".show(",
        ".legend(",
        ".grid(",
    )
    code_like_signals += sum(1 for line in content_lines if any(token in line for token in strong_substrings))

    python_statement_lines = sum(
        1
        for line in content_lines
        if re.match(r"^\s*(def|for|while|if|elif|else|return|with|try|except|class)\b", line)
    )
    code_like_signals += python_statement_lines

    indented_lines = sum(1 for line in content_lines if re.match(r"^\s{4,}\S", line))
    if indented_lines >= 2:
        code_like_signals += 1

    colon_block_lines = sum(1 for index, line in enumerate(content_lines[:-1]) if line.rstrip().endswith(":"))
    if colon_block_lines and indented_lines:
        code_like_signals += 1

    assignment_call_lines = sum(
        1
        for line in content_lines
        if "=" in line
        and "==" not in line
        and not line.lstrip().startswith(("#", "-", "*"))
        and re.search(r"=\s*[A-Za-z_][A-Za-z0-9_\.]*\s*\(", line)
    )
    if assignment_call_lines >= 1:
        code_like_signals += assignment_call_lines

    bracket_expression_lines = sum(
        1
        for line in content_lines
        if re.match(r"^\s*[A-Za-z_][A-Za-z0-9_]*(\[[^\]]+\])?\s*=", line)
        or re.search(r"\[[\"'][^\"']+[\"']\]", line)
    )
    if bracket_expression_lines >= 2:
        code_like_signals += 1

    return code_like_signals >= 2


def extract_markdown_cells_from_notebook(notebook_path: Path | str) -> list[str]:
    """Return markdown cell sources from a notebook in reading order."""
    try:
        import nbformat
    except ImportError as exc:
        raise ImportError(
            "nbformat is required for Markdown-to-Word rendering. Install it before running this notebook export section."
        ) from exc

    path = Path(notebook_path)
    if not path.exists():
        raise FileNotFoundError(f"Notebook file not found: {path}")
    notebook = nbformat.read(path, as_version=4)
    markdown_cells: list[str] = []
    for cell_index, cell in enumerate(notebook.cells):
        if getattr(cell, "cell_type", "") == "markdown":
            markdown_text = str(getattr(cell, "source", ""))
            if looks_like_python_code(markdown_text):
                print(
                    f"Warning: skipped suspicious markdown cell {cell_index} in {path.name} because it looks like Python code."
                )
                continue
            markdown_cells.append(markdown_text)
    return markdown_cells


def add_markdown_cells_to_docx(doc: Any, markdown_cells: list[str], run_dir: Path | str) -> None:
    """Render notebook markdown cells into a Word document."""
    if not markdown_cells:
        doc.add_paragraph("No markdown narrative was found in the notebook.")
        return

    state = SimpleNamespace(run_dir=Path(run_dir), equation_counter=0)
    for markdown_text in markdown_cells:
        render_markdown_text_to_docx(doc, markdown_text, state)


def get_markdown_cell_by_heading(markdown_cells: list[str], heading: str) -> str:
    """Return the first markdown cell containing the exact heading line."""
    normalized_heading = (heading or "").strip()
    if not normalized_heading:
        raise ValueError("heading must be a non-empty string.")

    for markdown_text in markdown_cells:
        lines = markdown_text.replace("\r\n", "\n").split("\n")
        if any(line.strip() == normalized_heading for line in lines):
            return markdown_text

    raise ValueError(f"Could not find notebook markdown section with heading: {heading}")


def strip_leading_h1_heading(markdown_text: str) -> str:
    """Remove a leading H1 heading when the report already has its own title."""
    lines = markdown_text.replace("\r\n", "\n").split("\n")
    if not lines:
        return markdown_text

    if re.match(r"^#\s+", lines[0].strip()):
        index = 1
        while index < len(lines) and not lines[index].strip():
            index += 1
        return "\n".join(lines[index:])

    return markdown_text


def dataframe_to_markdown_table(df: pd.DataFrame) -> str:
    """Render a small DataFrame as a simple Markdown table without extra dependencies."""
    if not isinstance(df, pd.DataFrame):
        raise TypeError("df must be a pandas DataFrame.")
    headers = [str(column) for column in df.columns]
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for record in df.to_dict(orient="records"):
        row_values = [str(record[column]) for column in headers]
        lines.append("| " + " | ".join(row_values) + " |")
    return "\n".join(lines)


def _is_table_line(line: str) -> bool:
    """Return True if a stripped line looks like a Markdown table row."""
    return line.startswith("|") and line.endswith("|") and line.count("|") >= 2


def _is_separator_line(line: str) -> bool:
    """Return True if a stripped line is a Markdown table separator (|---|---|)."""
    return bool(re.match(r"^\|[\s:]*-{2,}[\s:]*(\|[\s:]*-{2,}[\s:]*)*\|$", line))


def _parse_table_row(line: str) -> list[str]:
    """Split a Markdown table row into cell texts."""
    inner = line.strip("|")
    return [cell.strip() for cell in inner.split("|")]


def _flush_table_to_docx(doc: Any, table_rows: list[list[str]], has_header: bool) -> None:
    """Render accumulated table rows into a Word table."""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    if not table_rows:
        return

    n_cols = max(len(row) for row in table_rows)
    # Pad rows that have fewer columns
    for row in table_rows:
        while len(row) < n_cols:
            row.append("")

    table = doc.add_table(rows=len(table_rows), cols=n_cols)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(table_rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = _strip_inline_markdown(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Bold the header row
    if has_header and table_rows:
        for col_idx in range(n_cols):
            cell = table.cell(0, col_idx)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    doc.add_paragraph("")  # spacing after table


def render_markdown_text_to_docx(doc: Any, markdown_text: str, state: Any) -> None:
    """Render lightweight Markdown into a Word document."""
    from docx.shared import Inches

    text = (markdown_text or "").replace("\r\n", "\n")
    if not text.strip():
        return

    in_code_block = False
    in_display_math = False
    in_table = False
    code_buffer: list[str] = []
    display_math_buffer: list[str] = []
    paragraph_buffer: list[str] = []
    table_rows: list[list[str]] = []
    table_has_header = False

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        paragraph = " ".join(line.strip() for line in paragraph_buffer if line.strip())
        paragraph_buffer.clear()
        if paragraph:
            doc.add_paragraph(_strip_inline_markdown(paragraph))

    def flush_code_block() -> None:
        if not code_buffer:
            return
        for code_line in code_buffer:
            doc.add_paragraph(code_line, style="No Spacing")
        code_buffer.clear()

    def flush_display_math() -> None:
        if not display_math_buffer:
            return
        equation_text = "\n".join(display_math_buffer).strip()
        display_math_buffer.clear()
        if not equation_text:
            return
        equation_path = _render_display_equation_to_png(equation_text, state.run_dir, state.equation_counter)
        state.equation_counter += 1
        if equation_path is None:
            doc.add_paragraph(_latex_to_readable_text(equation_text))
            return
        try:
            doc.add_picture(str(equation_path), width=Inches(5.75))
        except Exception:
            doc.add_paragraph(_latex_to_readable_text(equation_text))

    def flush_table() -> None:
        nonlocal in_table, table_has_header
        if table_rows:
            _flush_table_to_docx(doc, table_rows, table_has_header)
            table_rows.clear()
        in_table = False
        table_has_header = False

    for raw_line in text.split("\n"):
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            flush_table()
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        if stripped == "$$" and not in_display_math:
            flush_paragraph()
            flush_table()
            in_display_math = True
            continue

        if stripped == "$$" and in_display_math:
            flush_display_math()
            in_display_math = False
            continue

        if not in_display_math and stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            flush_paragraph()
            flush_table()
            display_math_buffer.append(stripped[2:-2].strip())
            flush_display_math()
            continue

        if in_display_math:
            display_math_buffer.append(line)
            continue

        # Table handling
        if _is_table_line(stripped):
            if _is_separator_line(stripped):
                # Separator line — mark that the preceding row was a header
                if table_rows:
                    table_has_header = True
                continue
            flush_paragraph()
            if not in_table:
                in_table = True
            table_rows.append(_parse_table_row(stripped))
            continue

        # If we were in a table but this line is not a table line, flush it
        if in_table:
            flush_table()

        if not stripped:
            flush_paragraph()
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            heading_text = _strip_inline_markdown(heading_match.group(2).strip())
            if heading_text:
                doc.add_heading(heading_text, level=level)
            continue

        bullet_match = re.match(r"^[-*+]\s+(.*)$", stripped)
        if bullet_match:
            flush_paragraph()
            doc.add_paragraph(_strip_inline_markdown(bullet_match.group(1).strip()), style="List Bullet")
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered_match:
            flush_paragraph()
            doc.add_paragraph(_strip_inline_markdown(numbered_match.group(1).strip()), style="List Number")
            continue

        paragraph_buffer.append(line)

    flush_paragraph()
    flush_table()
    if in_code_block:
        flush_code_block()
    if in_display_math:
        flush_display_math()


def _add_single_content_section_to_docx(
    document: Any,
    section: dict[str, Any],
    run_dir: Path,
    state: Any,
) -> None:
    """Render one structured content section into a Word document."""
    from docx.shared import Inches

    heading = section.get("heading", "")
    body = section.get("body", "")
    figures = section.get("figures", [])

    if heading:
        document.add_heading(heading, level=1)

    if body:
        render_markdown_text_to_docx(document, body, state)

    for fig_entry in figures:
        if isinstance(fig_entry, dict):
            fig_name = fig_entry.get("filename", "")
            caption = fig_entry.get("caption", "")
        else:
            fig_name = str(fig_entry)
            caption = ""

        fig_path = run_dir / fig_name
        if not fig_path.exists():
            document.add_paragraph(f"Figure not found: {fig_name}")
            continue
        try:
            document.add_picture(str(fig_path), width=Inches(6.25))
            if caption:
                document.add_paragraph(caption)
        except Exception:
            document.add_paragraph(f"Could not embed figure: {fig_name}")


def _add_content_sections_to_docx(
    document: Any,
    content_sections: list[dict[str, Any]],
    run_dir: Path | str,
) -> None:
    """Render structured content sections into a Word document.

    Each section dict may contain:
    - ``heading``  (str)  – section title rendered as Heading 1
    - ``body``     (str)  – markdown text rendered into the document
    - ``figures``  (list) – figure filenames (str) or dicts with
      ``filename`` and optional ``caption`` keys, embedded inline
    """
    run_path = Path(run_dir)
    state = SimpleNamespace(run_dir=run_path, equation_counter=500)
    for section in content_sections:
        _add_single_content_section_to_docx(document, section, run_path, state)


def _add_exported_artifacts_to_docx(
    document: Any,
    csv_filenames: list[str] | tuple[str, ...] | None,
    figure_filenames: list[str] | tuple[str, ...] | None,
) -> None:
    """Render exported artifact lists into a Word document appendix."""
    document.add_heading("Exported artifacts", level=1)

    document.add_paragraph("CSV files")
    csv_names = [name for name in (csv_filenames or []) if name]
    if csv_names:
        for filename in csv_names:
            document.add_paragraph(filename, style="List Bullet")
    else:
        document.add_paragraph("No CSV files were saved.")

    document.add_paragraph("Figure files")
    figure_names = [name for name in (figure_filenames or []) if name]
    if figure_names:
        for filename in figure_names:
            document.add_paragraph(filename, style="List Bullet")
    else:
        document.add_paragraph("No figure files were saved.")


def build_word_report(
    report_title: str,
    setup_name: str,
    timestamp: str,
    run_dir: Path | str,
    notebook_path: Path | str,
    purpose_text: str,
    assumptions_text: str,
    parameter_summary: dict[str, Any] | list[Any] | tuple[Any, ...] | str | None,
    csv_filenames: list[str] | tuple[str, ...] | None,
    figure_filenames: list[str] | tuple[str, ...] | None,
    output_filename: str,
    content_sections: list[dict[str, Any]] | None = None,
) -> Path:
    """Build a Word report with notebook markdown narrative and embedded figures."""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for Word export. Install it before running this notebook export section."
        ) from exc

    run_path = Path(run_dir)
    notebook_file = Path(notebook_path)
    output_path = run_path / output_filename
    document = Document()

    document.add_heading(report_title, level=0)

    document.add_heading("Run Metadata", level=1)
    metadata_lines = [
        f"Setup name: {setup_name}",
        f"Run timestamp: {timestamp}",
        f"Notebook filename: {notebook_file.name}",
        f"Run directory: {run_path}",
    ]
    for line in metadata_lines:
        document.add_paragraph(line)

    document.add_heading("Notebook Narrative", level=1)
    markdown_cells = extract_markdown_cells_from_notebook(notebook_file)
    add_markdown_cells_to_docx(document, markdown_cells, run_path)

    document.add_heading("Purpose", level=1)
    document.add_paragraph(purpose_text or "No summary text provided.")

    document.add_heading("Assumptions", level=1)
    document.add_paragraph(assumptions_text or "No assumptions text provided.")

    document.add_heading("Parameter Summary", level=1)
    _add_parameter_summary(document, parameter_summary)

    if content_sections:
        _add_content_sections_to_docx(document, content_sections, run_path)

    document.add_heading("Exported CSV Files", level=1)
    csv_names = [name for name in (csv_filenames or []) if name]
    if csv_names:
        for filename in csv_names:
            document.add_paragraph(filename, style="List Bullet")
    else:
        document.add_paragraph("No CSV tables were saved.")
    document.add_paragraph("Detailed tables are saved separately as CSV files in the run folder.")

    has_inline_figures = any(section.get("figures") for section in (content_sections or []))

    document.add_heading("Figures", level=1)
    figure_names = [name for name in (figure_filenames or []) if name]
    if not figure_names:
        document.add_paragraph("No figure files were saved.")
    else:
        for filename in figure_names:
            figure_path = run_path / filename
            document.add_paragraph(filename, style="List Bullet")
            if has_inline_figures:
                continue
            if not figure_path.exists():
                document.add_paragraph(f"Skipped embedding {filename} because the file was not found.")
                continue
            try:
                document.add_picture(str(figure_path), width=Inches(6.25))
            except Exception:
                document.add_paragraph(f"Skipped embedding {filename} because it could not be inserted.")

    document.add_heading("Closing Note", level=1)
    document.add_paragraph(
        "This report captures the notebook narrative, summary metadata, and exported artifacts for this run."
    )

    document.save(str(output_path))
    return output_path


def build_structured_word_report(
    report_title: str,
    setup_name: str,
    timestamp: str,
    run_dir: Path | str,
    notebook_path: Path | str,
    output_filename: str,
    executive_summary_text: str | None = None,
    parameter_summary: dict[str, Any] | list[Any] | tuple[Any, ...] | str | None = None,
    report_sections: list[dict[str, Any]] | None = None,
    appendix_sections: list[dict[str, Any]] | None = None,
    csv_filenames: list[str] | tuple[str, ...] | None = None,
    figure_filenames: list[str] | tuple[str, ...] | None = None,
) -> Path:
    """Build a manuscript-style Word report from selected notebook sections and curated content."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for Word export. Install it before running this notebook export section."
        ) from exc

    run_path = Path(run_dir)
    notebook_file = Path(notebook_path)
    output_path = run_path / output_filename
    document = Document()

    document.add_heading(report_title, level=0)
    document.add_paragraph(f"Setup: {setup_name}")
    document.add_paragraph(f"Report timestamp: {timestamp}")
    document.add_paragraph(f"Source notebook: {notebook_file.name}")

    if executive_summary_text:
        document.add_heading("Executive summary", level=1)
        document.add_paragraph(executive_summary_text)

    if parameter_summary is not None:
        document.add_heading("Key parameters", level=1)
        _add_parameter_summary(document, parameter_summary)

    markdown_cells = extract_markdown_cells_from_notebook(notebook_file)
    state = SimpleNamespace(run_dir=run_path, equation_counter=1000)

    for section in report_sections or []:
        source = section.get("source", "content")
        if source == "notebook":
            markdown_text = get_markdown_cell_by_heading(markdown_cells, section.get("heading", ""))
            if section.get("strip_leading_h1", True):
                markdown_text = strip_leading_h1_heading(markdown_text)
            render_markdown_text_to_docx(document, markdown_text, state)
            continue
        if source == "content":
            _add_single_content_section_to_docx(document, section, run_path, state)
            continue
        raise ValueError(f"Unsupported report section source: {source!r}")

    if appendix_sections or csv_filenames or figure_filenames:
        document.add_heading("Appendices", level=1)

    for section in appendix_sections or []:
        source = section.get("source", "content")
        if source == "notebook":
            markdown_text = get_markdown_cell_by_heading(markdown_cells, section.get("heading", ""))
            if section.get("strip_leading_h1", True):
                markdown_text = strip_leading_h1_heading(markdown_text)
            render_markdown_text_to_docx(document, markdown_text, state)
            continue
        if source == "content":
            _add_single_content_section_to_docx(document, section, run_path, state)
            continue
        raise ValueError(f"Unsupported appendix section source: {source!r}")

    if csv_filenames or figure_filenames:
        _add_exported_artifacts_to_docx(document, csv_filenames, figure_filenames)

    document.save(str(output_path))
    return output_path


def build_structured_word_report_pandoc(
    report_title: str,
    setup_name: str,
    timestamp: str,
    run_dir: Path | str,
    notebook_path: Path | str,
    output_filename: str,
    executive_summary_text: str | None = None,
    parameter_summary: dict[str, Any] | list[Any] | tuple[Any, ...] | str | None = None,
    report_sections: list[dict[str, Any]] | None = None,
    appendix_sections: list[dict[str, Any]] | None = None,
    csv_filenames: list[str] | tuple[str, ...] | None = None,
    figure_filenames: list[str] | tuple[str, ...] | None = None,
    markdown_filename: str | None = None,
    reference_docx: Path | str | None = None,
    pandoc_path: str = "pandoc",
) -> Path:
    """Build a Word report via Pandoc so LaTeX math becomes native Word equations.

    The function writes an intermediate Markdown file into ``run_dir`` and then
    calls Pandoc to create ``output_filename``. Display and inline equations are
    preserved as TeX math in the Markdown so Pandoc can convert them to Word's
    native OMML equation format.
    """
    run_path = Path(run_dir)
    notebook_file = Path(notebook_path)
    output_path = run_path / output_filename
    markdown_path = run_path / (markdown_filename or f"{output_path.stem}.md")

    report_markdown = _build_structured_report_markdown(
        report_title=report_title,
        setup_name=setup_name,
        timestamp=timestamp,
        run_dir=run_path,
        notebook_path=notebook_file,
        executive_summary_text=executive_summary_text,
        parameter_summary=parameter_summary,
        report_sections=report_sections,
        appendix_sections=appendix_sections,
        csv_filenames=csv_filenames,
        figure_filenames=figure_filenames,
    )
    markdown_path.write_text(report_markdown, encoding="utf-8")

    command = [
        pandoc_path,
        str(markdown_path),
        "--from=markdown+tex_math_dollars+pipe_tables",
        "--to=docx",
        f"--output={output_path}",
        f"--resource-path={run_path}",
    ]
    if reference_docx is not None:
        command.append(f"--reference-doc={Path(reference_docx)}")

    try:
        completed = subprocess.run(command, capture_output=True, text=True, check=False)
    except FileNotFoundError as exc:
        raise FileNotFoundError(
            "Pandoc was not found. Install Pandoc or pass pandoc_path to build_structured_word_report_pandoc."
        ) from exc

    if completed.returncode != 0:
        raise RuntimeError(
            "Pandoc Word report export failed.\n"
            f"Command: {' '.join(command)}\n"
            f"stdout:\n{completed.stdout}\n"
            f"stderr:\n{completed.stderr}"
        )

    return output_path


def _build_structured_report_markdown(
    report_title: str,
    setup_name: str,
    timestamp: str,
    run_dir: Path,
    notebook_path: Path,
    executive_summary_text: str | None,
    parameter_summary: dict[str, Any] | list[Any] | tuple[Any, ...] | str | None,
    report_sections: list[dict[str, Any]] | None,
    appendix_sections: list[dict[str, Any]] | None,
    csv_filenames: list[str] | tuple[str, ...] | None,
    figure_filenames: list[str] | tuple[str, ...] | None,
) -> str:
    """Assemble the structured report content as Pandoc-friendly Markdown."""
    markdown_cells = extract_markdown_cells_from_notebook(notebook_path)
    chunks: list[str] = [
        f"# {report_title}",
        "",
        f"Setup: {setup_name}",
        "",
        f"Report timestamp: {timestamp}",
        "",
        f"Source notebook: {notebook_path.name}",
        "",
    ]

    if executive_summary_text:
        chunks.extend(["## Executive summary", "", executive_summary_text.strip(), ""])

    if parameter_summary is not None:
        chunks.extend(["## Key parameters", "", _parameter_summary_to_markdown(parameter_summary), ""])

    for section in report_sections or []:
        chunks.append(_section_to_markdown(section, markdown_cells, run_dir))

    if appendix_sections or csv_filenames or figure_filenames:
        chunks.extend(["## Appendices", ""])

    for section in appendix_sections or []:
        chunks.append(_section_to_markdown(section, markdown_cells, run_dir))

    if csv_filenames or figure_filenames:
        chunks.append(_exported_artifacts_to_markdown(csv_filenames, figure_filenames))

    return "\n".join(chunk.rstrip() for chunk in chunks if chunk is not None).strip() + "\n"


def _section_to_markdown(section: dict[str, Any], markdown_cells: list[str], run_dir: Path) -> str:
    """Render a structured report section as Markdown."""
    source = section.get("source", "content")
    if source == "notebook":
        markdown_text = get_markdown_cell_by_heading(markdown_cells, section.get("heading", ""))
        if section.get("strip_leading_h1", True):
            markdown_text = strip_leading_h1_heading(markdown_text)
        return markdown_text.strip() + "\n"
    if source == "content":
        return _content_section_to_markdown(section, run_dir)
    raise ValueError(f"Unsupported report section source: {source!r}")


def _content_section_to_markdown(section: dict[str, Any], run_dir: Path) -> str:
    """Render a curated content section as Markdown."""
    chunks: list[str] = []
    heading = str(section.get("heading", "") or "").strip()
    body = str(section.get("body", "") or "").strip()
    figures = section.get("figures", [])

    if heading:
        chunks.extend([f"## {heading}", ""])
    if body:
        chunks.extend([body, ""])

    for fig_entry in figures:
        if isinstance(fig_entry, dict):
            fig_name = str(fig_entry.get("filename", "") or "").strip()
            caption = str(fig_entry.get("caption", "") or fig_name).strip()
        else:
            fig_name = str(fig_entry).strip()
            caption = fig_name
        if not fig_name:
            continue
        fig_path = run_dir / fig_name
        if fig_path.exists():
            chunks.extend([f"![{_escape_markdown_image_caption(caption)}]({fig_name})", ""])
        else:
            chunks.extend([f"Figure not found: {fig_name}", ""])

    return "\n".join(chunks).strip() + "\n"


def _parameter_summary_to_markdown(parameter_summary: Any) -> str:
    """Render parameter summary content as Markdown."""
    if parameter_summary is None:
        return "No parameter summary provided."
    if isinstance(parameter_summary, dict):
        if not parameter_summary:
            return "No parameter summary provided."
        return "\n".join(f"- {key}: {_stringify(value)}" for key, value in parameter_summary.items())
    if isinstance(parameter_summary, (list, tuple)):
        if not parameter_summary:
            return "No parameter summary provided."
        lines: list[str] = []
        for item in parameter_summary:
            if isinstance(item, dict) and len(item) == 1:
                key, value = next(iter(item.items()))
                lines.append(f"- {key}: {_stringify(value)}")
            else:
                lines.append(f"- {_stringify(item)}")
        return "\n".join(lines)
    return _stringify(parameter_summary)


def _exported_artifacts_to_markdown(
    csv_filenames: list[str] | tuple[str, ...] | None,
    figure_filenames: list[str] | tuple[str, ...] | None,
) -> str:
    """Render exported artifact lists as Markdown."""
    chunks = ["## Exported artifacts", "", "CSV files", ""]
    csv_names = [name for name in (csv_filenames or []) if name]
    if csv_names:
        chunks.extend(f"- {name}" for name in csv_names)
    else:
        chunks.append("No CSV files were saved.")

    chunks.extend(["", "Figure files", ""])
    figure_names = [name for name in (figure_filenames or []) if name]
    if figure_names:
        chunks.extend(f"- {name}" for name in figure_names)
    else:
        chunks.append("No figure files were saved.")

    return "\n".join(chunks).strip() + "\n"


def _escape_markdown_image_caption(caption: str) -> str:
    """Escape characters that would break Markdown image alt text."""
    return caption.replace("[", "\\[").replace("]", "\\]")


def _add_parameter_summary(document: Any, parameter_summary: Any) -> None:
    """Render parameter summary content into a docx document."""
    if parameter_summary is None:
        document.add_paragraph("No parameter summary provided.")
        return
    if isinstance(parameter_summary, dict):
        if not parameter_summary:
            document.add_paragraph("No parameter summary provided.")
            return
        for key, value in parameter_summary.items():
            document.add_paragraph(f"{key}: {_stringify(value)}", style="List Bullet")
        return
    if isinstance(parameter_summary, (list, tuple)):
        if not parameter_summary:
            document.add_paragraph("No parameter summary provided.")
            return
        for item in parameter_summary:
            if isinstance(item, dict) and len(item) == 1:
                key, value = next(iter(item.items()))
                document.add_paragraph(f"{key}: {_stringify(value)}", style="List Bullet")
            else:
                document.add_paragraph(_stringify(item), style="List Bullet")
        return
    document.add_paragraph(_stringify(parameter_summary))


def _strip_inline_markdown(text: str) -> str:
    """Convert simple inline Markdown to plain readable text."""
    cleaned = _replace_inline_math(text)
    cleaned = re.sub(r"`([^`]*)`", r"\1", cleaned)
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"\*([^*]+)\*", r"\1", cleaned)
    cleaned = re.sub(r"__([^_]+)__", r"\1", cleaned)
    cleaned = re.sub(r"_([^_]+)_", r"\1", cleaned)
    cleaned = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", cleaned)
    cleaned = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", cleaned)
    return cleaned.strip()


def _replace_inline_math(text: str) -> str:
    """Replace inline LaTeX math with a readable plain-text approximation."""
    def repl(match: re.Match[str]) -> str:
        return _latex_to_readable_text(match.group(1))

    return re.sub(r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)", repl, text)


def _latex_to_readable_text(text: str) -> str:
    """Convert a subset of LaTeX math into readable plain text."""
    cleaned = text.strip()
    if not cleaned:
        return ""

    replacements = {
        r"\cdot": " * ",
        r"\times": " x ",
        r"\pm": " +/- ",
        r"\approx": " approx ",
        r"\leq": " <= ",
        r"\geq": " >= ",
        r"\neq": " != ",
        r"\to": " -> ",
        r"\rightarrow": " -> ",
        r"\left": "",
        r"\right": "",
        r"\,": " ",
        r"\;": " ",
        r"\:": " ",
        r"\!": "",
    }
    greek_map = {
        "alpha": "alpha",
        "beta": "beta",
        "gamma": "gamma",
        "delta": "delta",
        "epsilon": "epsilon",
        "eta": "eta",
        "theta": "theta",
        "lambda": "lambda",
        "mu": "mu",
        "nu": "nu",
        "pi": "pi",
        "rho": "rho",
        "sigma": "sigma",
        "tau": "tau",
        "phi": "phi",
        "omega": "omega",
        "Delta": "Delta",
        "Gamma": "Gamma",
        "Lambda": "Lambda",
        "Pi": "Pi",
        "Sigma": "Sigma",
        "Phi": "Phi",
        "Omega": "Omega",
    }

    for old, new in replacements.items():
        cleaned = cleaned.replace(old, new)

    cleaned = re.sub(r"\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}", r"(\1)/(\2)", cleaned)
    cleaned = re.sub(r"\\sqrt\s*\{([^{}]+)\}", r"sqrt(\1)", cleaned)
    cleaned = re.sub(r"\\(?:mathrm|text|operatorname)\s*\{([^{}]+)\}", r"\1", cleaned)
    cleaned = re.sub(r"_\{([^{}]+)\}", r"_\1", cleaned)
    cleaned = re.sub(r"\^\{([^{}]+)\}", r"^\1", cleaned)

    for name, replacement in greek_map.items():
        cleaned = cleaned.replace(f"\\{name}", replacement)

    cleaned = re.sub(r"\\([A-Za-z]+)", r"\1", cleaned)
    cleaned = cleaned.replace("{", "").replace("}", "")
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def _render_display_equation_to_png(equation_text: str, run_dir: Path, equation_index: int) -> Path | None:
    """Render a display equation to a PNG using matplotlib mathtext."""
    normalized = " ".join(line.strip() for line in equation_text.splitlines() if line.strip())
    if not normalized:
        return None

    output_path = run_dir / f"equation_{equation_index:03d}.png"
    figure: Figure | None = None
    try:
        figure = plt.figure(figsize=(0.01, 0.01))
        figure.patch.set_alpha(0.0)
        text_artist = figure.text(0.0, 0.5, f"${normalized}$", fontsize=16, ha="left", va="center")
        figure.canvas.draw()
        bbox = text_artist.get_window_extent(renderer=figure.canvas.get_renderer()).expanded(1.08, 1.25)
        width_in = max(bbox.width / figure.dpi, 0.5)
        height_in = max(bbox.height / figure.dpi, 0.35)
        figure.set_size_inches(width_in, height_in)
        text_artist.set_position((0.02, 0.5))
        figure.savefig(output_path, dpi=300, bbox_inches="tight", transparent=True, pad_inches=0.05)
        return output_path
    except Exception:
        output_path.unlink(missing_ok=True)
        return None
    finally:
        if figure is not None:
            plt.close(figure)


def _json_default(value: Any) -> Any:
    """Convert non-JSON-native values into serializable forms."""
    if isinstance(value, Path):
        return str(value)
    if hasattr(value, "item"):
        try:
            return value.item()
        except Exception:
            pass
    return str(value)


def _stringify(value: Any) -> str:
    """Convert values into concise report-friendly strings."""
    if isinstance(value, float):
        return f"{value:.6g}"
    if isinstance(value, (list, tuple)):
        return ", ".join(_stringify(item) for item in value)
    return str(value)
